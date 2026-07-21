"""
Stage 1 of the pipeline: read the regulator document and split it into the
individual OFAC observations. Deterministic, no LLM.

The input is irregular prose. Each observation starts with a header whose label
varies across the document — five styles seen in the sample, rotating:
    "Obs. 1 —", "Observación 2:", "Ítem 3.", "Caso 4 —", "Observación N.º 5."
We locate every header with one flexible regex and cut the text between
consecutive headers. The classifier/extractor never see the whole document,
only one observation at a time.
"""
from __future__ import annotations

import re
from pathlib import Path

# A header is an observation label + its case number. The label varies, so we
# match any of the known styles, optionally followed by "N.º". Group 1 = number.
_HEADER = re.compile(
    r"(?:Obs\.|Observaci[oó]n|[IÍ]tem|Caso)\s*(?:N\.?º\s*)?(\d{1,3})",
    re.IGNORECASE,
)


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF. pdfplumber handles the subset-embedded fonts
    these documents use, which trivial extractors mangle."""
    import pdfplumber  # imported lazily so a .docx-only run needn't have it

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    """Extract text from a Word document. The report's own wording ('un documento
    de Word') means a .docx is a first-class input, not just the sample PDF."""
    import docx  # imported lazily so a PDF-only run needn't have python-docx

    document = docx.Document(str(path))
    # Paragraph text plus any table cells — the report could arrive laid out either way.
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def load_text(path: str | Path) -> str:
    """Read the document to plain text, dispatching by file extension.

    Isolating the read behind this function keeps the rest of the pipeline
    independent of the input format: whatever the source, the segmenter and the
    agent only ever see plain text.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(
        f"Formato de documento no soportado: {suffix!r}. Se esperaba un archivo .pdf o .docx."
    )


def split_observations(text: str) -> list[tuple[int, str]]:
    """Cut the full text into (number, raw_text) per observation.

    Each observation runs from its own header up to the start of the next one
    (or the end of the document for the last one).
    """
    matches = list(_HEADER.finditer(text))
    observations: list[tuple[int, str]] = []
    for i, match in enumerate(matches):
        number = int(match.group(1))
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        raw = text[start:end].strip()
        observations.append((number, raw))
    return observations


def load_observations(path: str | Path) -> list[tuple[int, str]]:
    """Full stage 1: a document path -> list of (number, raw observation text)."""
    return split_observations(load_text(path))


if __name__ == "__main__":
    # Quick manual check: python src/loader.py [path]
    import sys

    sys.stdout.reconfigure(encoding="utf-8")  # Windows console: keep the accents
    src = sys.argv[1] if len(sys.argv) > 1 else "input/observaciones-ejemplo.pdf"

    obs = load_observations(src)
    print(f"Observations found: {len(obs)}")
    numbers = [n for n, _ in obs]
    print(f"Numbers: min={min(numbers)} max={max(numbers)} unique={len(set(numbers))}")
    for n, raw in obs[:2]:
        print(f"\n--- Obs {n} ({len(raw)} chars) ---\n{raw[:280]}...")
